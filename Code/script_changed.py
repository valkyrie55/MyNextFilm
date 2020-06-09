#import spacy

#total scenes = 13 ( including 1st line 'IT'S A MATCH')
import pandas as pd
import nltk
import re
from collections import Iterable
import random
import docx 
from docx.shared import Inches, Pt

random_list=[]
NUM_WORDS_INFILE=0
def random_gen():
    
    while(True):
        r=random.randint(1,NUM_WORDS_INFILE + 1000)
        if r not in random_list: 
            random_list.append(r)
            break
    return r
          
filename="itsamatch.txt"
characters = ['KUSH', 'POOJA', 'DRIVER', 'SHOPKEEPER', 'AKASH', 'RECEPTIONIST']
refined_line=[]

#word count
with open(filename, 'r') as f:   #open file 
    for line in f:
        words = line.split()
        NUM_WORDS_INFILE += len(words)   # count total no. of words
    print("Total number of words in the input file",NUM_WORDS_INFILE)

#read file
with open(filename, "r") as input:
    input_ = input.read().split('\n\n')

for line in input_:
    refined_line.append(line.strip())
#print(refined_line)

words=[]
speakers_words = [] #word list for speakers and their dialogues
scenes=[]
scene=[]
priority=[]
word_importance=[]
parenthetical='NONE'
dialogues=[]
actionline=[]
scene_dic = {}
refined_line=list(filter(lambda a: a != "", refined_line))
a_counter=0
# scene_number = 0               #added 28/May/2020
#dialogue_no = 0                #added 29/may
##CREATE SCENES LIST ALONG ACTION AND DIALOGUES IN THE GIVEN SEQUENCE   MAY UPDATE THE CODE FROM SHIVAM  ##RENU
for line in refined_line:            #if new scene
    if line.startswith('INT') or line.startswith('EXT') or line.startswith('EXT/INT') or line.startswith('INT/EXT'):
        a_counter=0
        #scene_number += 1
        dialogue_no = 0
        scenes.append(scene)
        #print(scene)
        scene=[]
        scene.append(line)
        scene_dic['SL']=line
        continue 
    else: #if the line is simple not a scene
        speakers_info = [[]]
        lis=line.split("\n")
        lis=[l.strip() for l in lis]
        word=lis[0]
        if word.split('(')[0].strip() in characters:
            mydic={}
            #spm=re.findall(r'\(([^()]+)\)',word)
            speaker=word.split('(')[0].strip()
            if len(lis)>1 and re.match(r"\(.*\)",lis[1]):
                parenthetical = lis[1]
                parenthetical=parenthetical.replace("\n","")
                dia=" ".join(lis[2:])
                dia=dia.replace("\n","")
            
            else:
                dia = " ".join(lis[1:])
                dia = dia.replace("\n","")
            if not (len(dia)==0 and parenthetical=="NONE"):
                list_speaker = []
                # dialogue_no += 1
                mydic[speaker]=[parenthetical,dia,len(dia)]
                dialogues.append(mydic)
                scene.append(mydic)
                #print("List of speakers in a scene: ", list_speaker)
                speakers_info.append(list_speaker)
                #print("speakers info: ", speakers_info)
            parenthetical = "NONE"
            #print("dialogue**********\n")
            scene_dic['D'] = mydic
            
        else:
            line = line.replace("\n","")
            line = ' '.join(line.split())
            actionline.append(line)
            scene.append(line)
            #print("action******\n")
            a_counter = a_counter+1

            scene_dic['A'+str(a_counter)] = line
    #print(scene_dic)
scenes.append(scene)
total_sc_cnt=1
for sc in scenes:
   # print(sc)
    total_sc_cnt+=1

print ("total scene count\n", total_sc_cnt)    
print("*************scene data********************\n")
#print ("SCENES LIST", scenes)

###  SCENE DATA IS READY, NOW CREATE THE WORD LIST ALONGWITH ITS ADDRESS TO WHICH IT BELONGS

new_scene=[]

scene_no=0
for scene in scenes:
    if scene_no == 0:
        scene_no = scene_no+1
        continue
    
    scene_dict={}
    action_counter=1
    dialogue_counter=1
    for line in scene:
        if type(line)==type(""):
            if line.startswith('INT') or line.startswith('EXT') or line.startswith('EXT/INT') or line.startswith('INT/EXT'):
                scene_dict['SL'] = line
                slug_words=line.split(" ")

                for each_word in slug_words:
                    temp_word={}
                    temp_word['word'] = each_word
                    temp_word['scene_num'] = scene_no
                    temp_word['type'] = 'SL'
                    temp_word['type_no'] = "" 
                    temp_word['sentence_no'] = '??'
                    temp_word['word_no_in sent'] = "??"
                    temp_word['POS'] = '??pos'
                    temp_word['importance'] = random_gen()
                    temp_word['zero_one'] = '1'
                    words.append(temp_word)

            else:
                scene_dict['AC'+ str(action_counter)] = line
                #print ("Scene_dic  ", scene_dict)
                sent_text = nltk.sent_tokenize(line)
                #print("Sent_text", sent_text)
                sentence_no=0
                for sentence in sent_text:
                    tokens = nltk.word_tokenize(sentence)
                    postag=nltk.pos_tag(tokens)
                    #print ("posttag", postag)
                    sentence_no=sentence_no+1
                    word_no=0
                    for (token,pos) in postag:
                        word_no=word_no+1
                        temp_word={}
                        temp_word['word'] = token
                        temp_word['scene_num'] = scene_no
                        temp_word['type'] = 'AC'
                        temp_word['type_no'] = str(action_counter)
                        temp_word['sentence_no'] = sentence_no
                        temp_word['word_no_in_sent'] = word_no
                        temp_word['POS'] = pos
                        temp_word['importance'] = random_gen()
                        temp_word['zero_one'] = '1'
                        words.append(temp_word)
                        #print("temp word in action")
                    #print("Words in action", action_counter, )
                    #print(words)
                action_counter = action_counter + 1

        else:
            dialogue_words="dialoge"
            diag_list = line.keys()
            diag_list=list(diag_list)
            par_dia=line.values()
            par_dia=list(par_dia)
            f=diag_list[0]+"###"+par_dia[0][0]+"###"+par_dia[0][1]
            #print(diag_list[0]+"###"+par_dia[0][0]+"###"+par_dia[0][1])
            scene_dict['DL'+ str(dialogue_counter)] = f   
            temp_word={}  #for word list
            temp_word['word'] = diag_list[0]
            temp_word['scene_num'] = scene_no
            temp_word['type'] = 'DL_SPEAKER'
            temp_word['type_no'] = str(dialogue_counter)  #dialogue counter in the particular scene
            temp_word['sentence_no'] = '??'
            temp_word['word_no_in_sent'] = "??"
            temp_word['POS'] = 'NNP'
            temp_word['importance'] = random_gen()
            temp_word['zero_one'] = '1'
            words.append(temp_word)
            #print(temp_word) 
            parentheticals=par_dia[0][0].split(" ")
            word_no=0
            
            for each_word in parentheticals:
                if each_word=='NONE':
                    break
                word_no= word_no+1
                each_word = each_word.replace("(", "")
                each_word = each_word.replace(")", "")
                temp_word={}
                temp_word['word'] = each_word
                temp_word['scene_num'] = scene_no
                temp_word['type'] = 'DL_PARENTH'
                temp_word['type_no'] = str(dialogue_counter)
                temp_word['sentence_no'] = 1
                temp_word['word_no_in_sent'] = word_no
                temp_word['POS'] = '??pos'
                temp_word['importance'] = random_gen()
                temp_word['zero_one'] = '1'
                words.append(temp_word)
                #print(temp_word) +
            dialogues=par_dia[0][1].split(" ")
            #print("dial=============", dialogues)
            sent_text = nltk.sent_tokenize(par_dia[0][1])
            sentence_no=0
            for sentence in sent_text:
                tokens = nltk.word_tokenize(sentence)
                postag=nltk.pos_tag(tokens)
                sentence_no=sentence_no+1
                word_no=0
                for (token,pos) in postag:
                    word_no=word_no+1
                    temp_word={}
                    temp_word['word'] = token
                    temp_word['scene_num'] = scene_no
                    temp_word['type'] = 'DL_DELIVERY'
                    temp_word['type_no'] = str(dialogue_counter)
                    temp_word['sentence_no'] = sentence_no
                    temp_word['word_no_in_sent'] = word_no
                    temp_word['POS'] = pos
                    temp_word['importance'] = random_gen()
                    temp_word['zero_one'] = '1'
                    words.append(temp_word)   
            dialogue_counter = dialogue_counter + 1
    ##see this
   
    #if scene_no > 1 : do we really need this check 
    new_scene.append(scene_dict)
    scene_no = scene_no+1
#print (" \n LINE 269 NEW SCENES ", new_scene)
# print(" WORDS LIST\n")
# for word in words:
#     print(word)
#     print("\n")
print (" LEN OF WORDLIST", len(words))

def find_threshold(reduced_percent):
    """ Takes as input the percentage of compression required in the script 
    and returns the 'threshold_val' - the thershold value of IMPORTANCE 
                the 'threshold_count' - the number of words in the reduced script
    """
    number_of_words = 0 
    word_temp = []
    for each_word in words:
        word_temp.append(each_word['importance'])
        number_of_words = number_of_words+1   
        #print (word_temp[number_of_words-1],number_of_words)   
    word_temp.sort() 
    reduced_percent=1-reduced_percent/100   #output script
    print("redu=",reduced_percent)
    threshold_count = round(number_of_words*reduced_percent)
    threshold_val = word_temp[threshold_count]
    print("LINE 292 threshold in fnc   value and count = ", threshold_val, threshold_count)
    return threshold_val, threshold_count

def find_threshold_removal_impact(reduced_percent):
    """ Takes as input the percentage of compression required in the script 
    and returns the 'threshold_val' - the thershold value of REMOVAL IMPACT 
                the 'threshold_count' - the number of words in the reduced script
    """
    number_of_words=0 
    word_temp=[]
    for each_word in words:
        word_temp.append(each_word['removal_impact'])
        number_of_words= number_of_words+1   
        #print (word_temp[number_of_words-1],number_of_words)   
    word_temp.sort() 
    reduced_percent=1-reduced_percent/100
    #print("redu=",reduced_percent)
    threshold_count = round(number_of_words*reduced_percent)
    threshold_val = word_temp[threshold_count]
    print("\n LINE 311 threshold removal_impact in fnc  value : \t and  count  = ", threshold_val, threshold_count)
    return threshold_val, threshold_count



def set_zero_initial(reduced_percent):
    """ Takes as input the percentage of reduced script and sets the words as '0' and '1' based on 
        the threshold value of importance

    """
    threshold_v, threshold_c = find_threshold(reduced_percent) #find threshold
    word_no = 1
    cnt=0
    non_rem_cnt=0
    print ("threshold value ", threshold_v)
    for each_word in words:
        if each_word['importance'] < threshold_v:
            each_word['zero_one'] = '0'     #remove
            cnt+=1
           
        else:
            each_word['zero_one'] = '1'
            non_rem_cnt+=1
        word_importance.append(each_word['importance'])
        #print("importance zero=", each_word['zero_one'])
    print (" Number of words below threshold value", cnt)
    print (" Number of words above threshold value", non_rem_cnt)

    return threshold_c


def set_zero_initial_removal_impact(reduced_percent):
    """ Takes as input the percentage of reduced script and sets the words as '0' and '1' based on 
        the threshold value of removal impact
    """
    threshold_v, threshold_c = find_threshold_removal_impact(reduced_percent) 
    word_no=1
    for each_word in words:
        #print("impact=",each_word['removal_impact'],"threshold=", threshold_v )
        if each_word['removal_impact'] <= threshold_v:
            each_word['zero_one'] = '0'  #removable  
        else:
            each_word['zero_one'] = '1'
        #print("impact zero=", each_word['zero_one'])
    return threshold_c


def convert_importance_to_priority(word_importance):
    """
    Takes the 'word_importance' as the input and define the 'priority' of the word for removal based on WORD IMPORTANCE
    e.g. if priority[1] = j   this means that the first priority for removal is of the  'jth WORD' 
            priority[5] = 7   means 5th priority for removal is of the 7th word
            CAN BE MODIFIED USING SORTING OF REDUCED WORD IMPORTANCE
    """
    reduced_word_imp = list(word_importance)
    length = len(word_importance)
    minpos = word_importance.index(min(word_importance))
    #minpos is the index of the min value in the word_importance list
    k = 0
    for i in range(length):
        priority.append(0)    #set all values as 0

    while(k < length):
        min_value = min(reduced_word_imp)          #reduced_word_imp is word_importance
        minpos = word_importance.index(min_value)
        priority[k] = minpos                      # first value in priority list will be the index of the min value in word_importance list
        reduced_word_imp.remove(min_value)        # that word is removed
        k = k + 1

def convert_removal_impact_to_priority(removal_impact_list):
    """
    Takes the 'removal impact list' as the input and define the 'priority' of the word for removal based on REMOVAL IMPACT VALUE OF THE WORD
    e.g. if priority[1]=j   this means that the first priority is of the  'jth WORD' (based on removal impact of word j)
            priority[5]=7   means that the 5th priority for removal is of the 7th word (based on removal impact of word 7)
    """
    unique_removal_impact_list = list(dict.fromkeys(removal_impact_list))
    unique_removal_impact_list.sort()
    length = len(removal_impact_list)
    priority_removal=[]
    for i in range(length):
        priority_removal.append(0)
    #print("small list",unique_removal_impact_list)
    priority_counter=0
    for each_unique in unique_removal_impact_list:
        word_no=0
        for each_removal_impact_value in removal_impact_list:
            if (each_unique == each_removal_impact_value):
                #print("uniq=",each_unique,  "all impact=",each_removal_impact_value)
                priority_removal[priority_counter] = word_no
                # if priority[1]=j   this means that the first priority is of the  'jth WORD' (based on removal impact of word j)
                # priority[5]=7
                priority_counter=priority_counter+1
            word_no=word_no+1
    #print("removal priority=",priority_removal)            
    return priority_removal


def removal_impact_sentence(word):
    """ Takes the word as input and calculates the removal impact, if the sentence it contains needs to be removed
        Values returned are
        'word_removal_impact' - sum of importance of all the words in the sentence that contains it
        'num_impacted'        - number of words in the sentence
        'impacted'            - list of word_no of the words in the the sentence
    """
    #print("word to check in sentence =",word['word'])
    comp_scene_no= word['scene_num']
    #print("comp=",comp_scene_no)
    comp_scene_type=word['type']
    comp_type_no=word['type_no']
    comp_sentence_no = word['sentence_no']
    word_no=0
    num_impacted=0
    impacted=[]
    word_removal_impact=0
    for each_word in words:
        #print("word_no=", word_no,"word is=", each_word)
        if(each_word['scene_num'] ==comp_scene_no and each_word['type'] == comp_scene_type and each_word['type_no'] == comp_type_no and each_word['sentence_no'] == comp_sentence_no):
            #print("word=", each_word['word'], "word_no=", word_no, "importance=",each_word['importance'] )
            word_removal_impact= word_removal_impact+ each_word['importance']
            impacted.append(word_no)
            num_impacted = num_impacted+1
        word_no=word_no+1
    #print("removal impact=", word_removal_impact)   
    #print("impacted list=", impacted) 
    return word_removal_impact, num_impacted, impacted      

def removal_impact_scene(word):
    """Takes the word as input and calculates the removal impact, if the SCENE it contains needs to be removed
        Values returned are
        'word_removal_impact' - sum of importance of all the words in the SCENE that contains it
        'num_impacted'        - number of words in the SCENE
        'impacted'            - list of word_no of the words in the SCENE
    """
    #print("removal of word in  scene=",word['word'])
    comp_scene_no = word['scene_num']    # compute scene number
    #print("comp=",comp_scene_no)
    impacted = []                 #list of word_no of the words in the SCENE              
    num_impacted = 0
    word_no=0
    scene_removal_impact=0
    for each_word in words:
        if(each_word['scene_num'] == comp_scene_no):   #loop will execute for all the words in that scene 
            #print("Scene  word=", each_word['word'], "scene_num=", comp_scene_no, "scene_no=",each_word['scene_num'] )
            scene_removal_impact = scene_removal_impact + each_word['importance']   #the scene removal impact will b the sum of the word importance of all the words in that scene
            #print(each_word)
            impacted.append(word_no)
            #print("word_no=", word_no)
            num_impacted = num_impacted+1
        word_no=word_no+1
    #print("scene removal impact=", scene_removal_impact) 
    #print("impacted list=", impacted)  
    return scene_removal_impact, num_impacted, impacted 



def removal_impact_speaker(word):
    """Takes the word as input and calculates the removal impact, if the word is the SPEAKER of a dialogue
        Values returned are
        'word_removal_impact' - sum of importance of all the words including the corresponding DIALOGUE, PARENTHEICAL and SPEAKER
        'num_impacted'        - number of words impacted
        'impacted'            - list of word_no of the words that are impacted
    """
    comp_scene_no= word['scene_num']
    comp_scene_type=word['type']
    comp_type_no=word['type_no']
    num_impacted=0
    impacted=[]
    word_no=0
    speaker_removal_impact=0
    for each_word in words:
        if(each_word['scene_num'] == comp_scene_no and each_word['type_no'] == comp_type_no):
            if (each_word['type'] == 'DL_DELIVERY' or each_word['type'] == 'DL_PARENTH' or each_word['type'] == 'DL_SPEAKER'):
               # print("Scene  word=", each_word['word'], "word_no=", word_no, "importance=",each_word['importance'] )
               # print("found", each_word['word'])
                speaker_removal_impact = speaker_removal_impact + each_word['importance']
                impacted.append(word_no)
                num_impacted = num_impacted + 1
        word_no = word_no + 1
    #print("speaker removal impact=", speaker_removal_impact) 
    #print("impacted list=", impacted)   
    return speaker_removal_impact, num_impacted, impacted

def assign_word_removal_impact():
    """
    Calculates the removal impact of each word in the script 
        If the word is part of a slugline, the removal impact is sum of importance of all the words in the scene
        If the word is a speaker, the removal impact is sum of importance of all words in the dialoge, parenthetical and speaker
        If the word is a verb, the removal impact is sum of all words in the sentence that contains it
        DEFAULT REMOVAL IMPACT CHANGED FROM ZERO TO 1
    """
    removal_impact_value_list=[]
    for each_word in words:
        impacted_words=[]   #list of words impacted by the removal of the each_word
        #print("  word=", each_word['word'])
        each_word['removal_impact'] = 1
        each_word['impacted_word_list'] = impacted_words

        if(each_word['type'] =='SL'):   #slugline
            removal_impact, impacted_no, impacted_words = removal_impact_scene(each_word)
            each_word['removal_impact'] = removal_impact
            each_word['impacted_word_list'] = impacted_words

        if (each_word['type'] == 'DL_SPEAKER'):    #same for dialogue
            removal_impact, impacted_no, impacted_words = removal_impact_speaker(each_word)
            each_word['removal_impact'] = removal_impact
            each_word['impacted_word_list'] = impacted_words

        str=each_word['POS']
        #print (str, str[0:2])
        if (str[0:2] == 'VB'):
            removal_impact, impacted_no, impacted_words=removal_impact_sentence(each_word)
            each_word['removal_impact'] = removal_impact
            each_word['impacted_word_list'] = impacted_words
        
        #print ("removal impact =", each_word['removal_impact'])
        removal_impact_value_list.append(each_word['removal_impact'])
    return removal_impact_value_list


def calculate_a_by_b():
    """
    Calculate the theshold value of A/B that decides the stopping criteria
    """
    total_importance = 0
    zero_importance = 0
    for each_word in words:
        #print("each=", each_word)
        if each_word['zero_one'] == '0': #removed
            zero_importance = zero_importance + each_word['importance']
        total_importance = total_importance + each_word['importance']  
    B = zero_importance / total_importance      
    #print("zero imp =", zero_importance, "total_imp = ", total_importance, "B=",B)
    A = (100-reduced_percent)/100
    #print("A = ", A)
    stop_threshold = A/B
    #stop the execution once 
    #print("a_by_b_threshold=", stop_threshold)
    return stop_threshold

print ("\n  WORDS HAVE BEEN ASSIGNED IMPORTANCE AND ADDRESS")
    
reduced_percent = 10
threshold_counter = set_zero_initial(reduced_percent)   
#got word_importance list
convert_importance_to_priority(word_importance)
#print("word importance=", word_importance)
print()
#print("importance priority=",priority)
g_removal_impact_value_list = assign_word_removal_impact()   #calculate word impact of each word
#print()
#print("removal impact values list=", g_removal_impact_value_list)
threshold_counter_removal_impact = set_zero_initial_removal_impact(reduced_percent)

r_priority = convert_removal_impact_to_priority(g_removal_impact_value_list)

a_by_b_threshold = calculate_a_by_b()

#print (" \n  REMOVAL IMPACT of each word")

# for each_word in words:
#     #print("word=",each_word)
#     print( "word=", each_word['word'], "impact=", each_word['removal_impact'], "zero_one=", each_word['zero_one'] )

# for each_word in words:
#     print("word=",each_word)
#     print( "word=", each_word['word'], "importance=", each_word['importance'])

#print ("\n list of words to be removed")
#print(r_priority)
print ("Len of r_priority: ",len(r_priority))
print("count of words to be removed: ", threshold_counter_removal_impact)
print("len of word list: ", len(words))
c = 0
for w in words:
    new_dict = {'zero_one': '1'}
    w.update(new_dict)
for w in words:
    if w['zero_one'] == '1':
        c +=1
print("Count of words whose zero_one value has been changed to 1: ",c)

count = 0
for i in r_priority[ : threshold_counter_removal_impact]:
    for each_word in words:
        if i == words.index(each_word):
            new_dict = {'zero_one': '0'}
            each_word.update(new_dict)              # to be removed
            count += 1
# print("Len of r_words: ", len(r_words))
print("Words whose zero_one values have been set to 0: ", count)

###########################################################################################33
#generate script

doc = docx.Document() 
style = doc.styles['Normal']
font = style.font
font.name = 'Courier New'
font.size = Pt(12)

def print_speaker(dialogue_speaker):
    spk = doc.add_paragraph(dialogue_speaker)
    spk_format = spk.paragraph_format
    spk_format.space_after = 0
    spk_format.line_spacing = Pt(12)
    spk_format.left_indent = Inches(2.5)

def print_dialogue(dialogue_delivery):
    d = doc.add_paragraph()
    d.add_run(dialogue_delivery)
    d_format = d.paragraph_format
    d_format.keep_together = True
    d_format.space_after = Pt(12)
    d_format.line_spacing = Pt(12)
    d_format.left_indent = Inches(1.5)
    d_format.right_indent = Inches(1.25)

def create_script(doc, dialogue_speaker, dialogue_delivery, dialogue_parenthetical, slugline_list, action_list): 
#slugline 
    slugline_data = " ".join(slugline_list)
    #print("slugline: ", slugline_data)
    slug_data = doc.add_paragraph()
    slug_data_format = slug_data.paragraph_format   #styling
    slug_data_format.space_after = Pt(12)
    slug_data_format.keep_with_next = True   #next line
    slug_data_format.left_indent = Inches(-0.5)
    slug_data.add_run(slugline_data).bold = True
 
#action line
    action_data = " ".join(action_list)
    #print("Action line: ", action_data)
    act = doc.add_paragraph()
    act.add_run(action_data)
    act_format = act.paragraph_format
    act_format.space_after = Pt(12)
    act_format.line_spacing = Pt(12)
    act_format.left_indent = Inches(0.5)

#speaker
    for i in range(len(dialogue_speaker)):
        print_speaker(dialogue_speaker[i])
        print_dialogue(dialogue_delivery[i])

    doc.save('output_10.docx')     

scene_no = 13
for s in range(1, scene_no+1):
    dialogue_counter = 1
    dialogue_speaker = []                       # speakers in the scene
    dialogue_delivery = []                      # list of words in the dialog of the particular scene
    dialogue_parenthetical = []                 # list of words in the parenthetical of the particular scene
    action_list = []                            # list of words in the action of the particular scene
    slugline_list = []                          # list of words in the slugline of the particular scene                               

    print("Scene taken is: ", s)
    cnt_total = 0
    cnt_removed = 0
    for w in words:
        if w['scene_num'] == s:
            cnt_total += 1
            if w['type'] == 'SL':                     #SLUGLINE
                if w['zero_one'] == '0':   #to be removed
                    cnt_removed += 1
                    res = ''
                    for t in w['word']:
                        res = res + t +  '\u0336'   #strike through
                    slugline_list.append(res)
                else:
                    slugline_list.append(w['word'])

            elif w['type'] == 'AC':   
                if w['zero_one'] == '0':   #to be removed
                    cnt_removed += 1
                    res = ''
                    for t in w['word']:
                        res = res + t +  '\u0336'   #strike through
                    action_list.append(res)
                else:
                    action_list.append(w['word'])

            elif w['type'] == 'DL_SPEAKER':   
                if w['zero_one'] == '0':   #to be removed
                    cnt_removed += 1
                    res = ''
                    for t in w['word']:
                        res = res + t +  '\u0336'   #strike through
                    dialogue_speaker.append(res)
                else:
                    dialogue_speaker.append(w['word'])  #list of all speakers
                dd = []
                for w in words:
                    if w['scene_num'] == s and w['type'] == 'DL_DELIVERY'and w['type_no'] == str(dialogue_counter): 
                        if w['zero_one'] == '0':   #to be removed
                            cnt_removed += 1
                            res = ''
                            for t in w['word']:
                                res = res + t +  '\u0336'   #strike through
                            dd.append(res)
                        else:
                            dd.append(w['word'])
                    # elif w['scene_num'] == s and w['type'] == 'DL_PARENTH'and w['type_no'] == str(dialogue_counter): 
                    #     if w['zero_one'] == '0':   #to be removed
                    #         res = ''
                    #         for t in w['word']:
                    #             res = res + t +  '\u0336'   #strike through
                    #         dd.append(res)
                    #     else:
                    #         dd.append(w['word'])
                val = ' '.join(dd)
                dialogue_delivery.append(val)
                dialogue_counter += 1
    print("Total words in ", s," scene: ", cnt_total)
    print("words removed from",s," scene : ", cnt_removed)
    create_script(doc ,dialogue_speaker, dialogue_delivery, dialogue_parenthetical, slugline_list, action_list)
    # print("---------------------------------------------")    
print("Done!")