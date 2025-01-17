# latest code
######################################################### PART I CODE ###############################################################
import re
import nltk
import spacy
import json
import textacy
import docx 
from docx.shared import Inches, Pt
from collections import Counter
from spacy.lang.en.stop_words import STOP_WORDS
from spacy.tokenizer import Tokenizer
from nltk.tokenize import word_tokenize
from nltk.tree import *
from string import punctuation
from nltk.tokenize import RegexpTokenizer
from nltk.stem import WordNetLemmatizer
from nltk.corpus import stopwords
STOP_WORDS = set(stopwords.words('english'))
nlp = spacy.load('en_core_web_sm')
def warn(*args, **kwargs):                      #to supress warnings
    pass
import warnings
warnings.warn = warn

random_list=[]

def create_character_list():
    for i in data['data_file']:
        characters.append(i['name'])

def preprocessed_sentence(sentence):                 # for removing punchuations
    t = nltk.RegexpTokenizer(r"\w+")
    new_sentence = t.tokenize(sentence)
    return ' '.join(new_sentence)

def filt(x):
    return x.label()=='NP'                           # for Noun Phrases

def find_noun_phrases(sentence):
  noun = []
  tokens = nltk.pos_tag(word_tokenize(sentence))
  pattern = " NP:{<DT>?<JJ>*<NN>}"                   # pattern for NP
  ch = nltk.RegexpParser(pattern)
  tree = ch.parse(tokens)
  for subtree in tree.subtrees(filter =  filt): 
    res = ''
    for i in range(0, len(subtree)):
      res = res + subtree[i][0] + ' '
    noun.append(res)
  return noun

def find_phrases(sentence, scene_no, sentence_no_in_scene, sent_type, paragraph_no):
    if sentence:
        verb = [] 
        noun = [] 
        sentence = sentence.lower()
        pattern = r'(<VERB>?<ADV>*<VERB>+)'
        docx = textacy.make_spacy_doc(sentence, lang = "en_core_web_sm")
        verb_phrases_list = textacy.extract.pos_regex_matches(docx, pattern)
        noun = find_noun_phrases(sentence)
        val = 1
        for noun_phrase in noun:
            if noun_phrase: 
                temp = {}
                temp['phrase'] = noun_phrase
                temp['importance'] = val
                temp['phrase_type'] = 'NP'
                # append phrase data in sentences list
                for each_sentence in sentences:
                    if each_sentence['type'] == 'AC':
                        if each_sentence['scene_no'] == scene_no and each_sentence['sentence_no_in_scene'] == sentence_no_in_scene:
                            each_sentence['phrases'].append(temp)                   
                temp['scene_no'] = scene_no
                temp['sentence_no_in_scene'] = sentence_no_in_scene
                temp['type'] = sent_type
                temp['paragraph_no'] = paragraph_no
                phrases.append(temp)
                val += 1
        # apply same procedure for verb phrases
        for chunk in verb_phrases_list:
            verb.append(chunk.text)  
        #append in the phrases list
        for verb_phrase in verb:
            if verb_phrase:
                temp = {}
                temp['phrase'] = verb_phrase
                temp['importance'] = val
                temp['phrase_type'] = 'VP'
                for each_sentence in sentences:
                    if each_sentence['type'] == 'AC':
                        if each_sentence['scene_no'] == scene_no and each_sentence['sentence_no_in_scene'] == sentence_no_in_scene:
                            each_sentence['phrases'].append(temp)
                temp['scene_no'] = scene_no
                temp['sentence_no_in_scene'] = sentence_no_in_scene
                temp['type'] = sent_type
                temp['paragraph_no'] = paragraph_no
                phrases.append(temp)
                val += 1


# extract characters from .json file
def create_character_list():
    for i in data['data_file']:
        characters.append(i['name'])

# extract character importance from .json file
def character_importance(character_name):
    for i in data['data_file']:
        if i['name'] == character_name.upper() :
            return i['character_importance']

def set_importance(phrase, sent_no, scene_no, importance):
    for each_phrase in phrases:
        if each_phrase['scene_no'] == scene_no and each_phrase['type'] == 'AC' and each_phrase['sentence_no_in_scene'] == sent_no and each_phrase['phrase'].lower() == phrase.lower():          
            each_phrase['importance'] = each_phrase['importance'] * importance 

def set_importance1(ph, sent_no, scene_no, importance):
    for each_phrase in phrases:
        if each_phrase['scene_no'] == scene_no and each_phrase['type'] == 'AC' and each_phrase['sentence_no_in_scene'] == sent_no and each_phrase['phrase'].lower() == ph.lower() :    
                each_phrase['importance'] = each_phrase['importance'] * importance 
#calculate imp of the sentences
def set_importance2(phrase, scene_no, sent_no):
    for each_phrase in phrases:
        if each_phrase['scene_no'] == scene_no and each_phrase['type'] == 'AC' and each_phrase['sentence_no_in_scene'] == sent_no and each_phrase['phrase'].lower() == phrase.lower() : 
            return each_phrase['importance']
    return 0

# read file
filename = "itsamatch.txt"
refined_line = []
try:
    with open(filename, "r") as input:
        input_ = input.read().split('\n\n')
    print(" File read successfully")
except: 
    print(" Error reading file!")

for line in input_:
    refined_line.append(line.strip())

#read json file
try:
    with open('/home/shweta/work/MNF_work/papers/vectors/itsamatchCharacterVector.json') as f:
        data = json.load(f)
except:
    print("Error in reading json file")
words = []
phrases = []
speakers_words = [] #word list for speakers and their dialogues
scenes=[]
characters = []
scene=[]
priority=[]
phrase_importance = []
sent_importance=[]
parenthetical='NONE'
dialogues=[]
actionline=[]
scene_dic = {}
t = ()
refined_line=list(filter(lambda a: a != "", refined_line))
a_counter=0

create_character_list() 

##CREATE SCENES LIST ALONG ACTION AND DIALOGUES IN THE GIVEN SEQUENCE   MAY UPDATE THE CODE FROM SHIVAM  ##RENU
for line in refined_line:            #if new scene
    if line.startswith('INT') or line.startswith('EXT') or line.startswith('EXT/INT') or line.startswith('INT/EXT'):
        a_counter=0
        #scene_number += 1
        dialogue_no = 0
        scenes.append(scene)
        #print(scene)
        scene=[]
        scene.append(line)
        scene_dic['SL']=line
        continue 
    else: #if the line is simple not a scene
        speakers_info = [[]]
        lis=line.split("\n")
        lis=[l.strip() for l in lis]
        word=lis[0]
        if word.split('(')[0].strip() in characters:
            mydic={}
            #spm=re.findall(r'\(([^()]+)\)',word)
            speaker=word.split('(')[0].strip()
            if len(lis)>1 and re.match(r"\(.*\)",lis[1]):
                parenthetical = lis[1]
                parenthetical=parenthetical.replace("\n"," ")
                dia = " ".join(lis[2:])
                dia = dia.replace("\n"," ")
            
            else:
                dia = " ".join(lis[1:])
                dia = dia.replace("\n"," ")
            if not (len(dia)==0 and parenthetical=="NONE"):
                list_speaker = []
                # dialogue_no += 1
                mydic[speaker]=[parenthetical,dia,len(dia)]
                dialogues.append(mydic)
                scene.append(mydic)
                speakers_info.append(list_speaker)
            parenthetical = "NONE"
            scene_dic['D'] = mydic
            
        else:
            line = line.replace("\n"," ")
            line = ' '.join(line.split())
            actionline.append(line)
            scene.append(line)
            #print("action******\n")
            a_counter = a_counter+1

            scene_dic['A'+str(a_counter)] = line
    #print(scene_dic)
scenes.append(scene)

# create sentences list
sentences = []
scene_no = 1
for each_scene in scenes[1: len(scenes)]:
    paragraph_no = 0
    sentence_counter = 1
    action_counter = 1
    for sentence in each_scene:
        if type(sentence) == type(""):                        
            if sentence.startswith('INT') or sentence.startswith('EXT') or sentence.startswith('EXT/INT') or sentence.startswith('INT/EXT'):
                temp = {}
                temp['sentence'] = sentence 
                temp['scene_no'] = scene_no
                temp['sentence_no_in_scene'] = 0
                temp['type'] = 'SL'
                temp['type_no'] = ''
                temp['speaker'] = '??'
                temp['phrases'] = []
                temp['importance'] = 0
                temp['final_importance'] = 0
                temp['zero_one'] = '1'
                temp['paragraph_no'] = paragraph_no
                sentences.append(temp)
            else:       
                paragraph_no += 1                                      # AC line
                sentence_list = sentence.split('.')                    # sentence list
                val = 1
                for sentence in sentence_list:  
                    if sentence:                  
                        temp = {}
                        temp['sentence'] = sentence + '. '
                        temp['scene_no'] = scene_no
                        temp['sentence_no_in_scene'] = sentence_counter
                        temp['type'] = 'AC'
                        temp['type_no'] = str(action_counter)
                        temp['speaker'] = '??'
                        temp['phrases'] = []
                        temp['final_importance'] = 0
                        temp['zero_one'] = '1'
                        temp['importance'] = val
                        temp['paragraph_no'] = paragraph_no
                        sentence_counter = sentence_counter + 1
                        sentences.append(temp)
                        action_counter += 1
                        val += 1
        else:                      # dialogues
            dialogue_words = "dialoge"
            diag_list = sentence.keys()
            diag_list = list(diag_list)
            par_dia = sentence.values()
            par_dia = list(par_dia)
            f = diag_list[0]+"###"+par_dia[0][0]+"###"+par_dia[0][1]
            # scene_dict['DL'+ str(dialogue_counter)] = f   
            temp = {} 
            temp['sentence'] = diag_list[0]          # speaker
            temp['scene_no'] = scene_no
            temp['type'] = 'DL_SPEAKER'
            temp['phrases'] = []
            temp['type_no'] = ''  #dialogue counter in the particular scene
            temp['sentence_no'] = 0
            temp['paragraph_no'] = 0
            temp['importance'] = 0
            temp['final_importance'] = 9999
            temp['zero_one'] = '1'
            sentences.append(temp)
            parentheticals = par_dia[0][0].split(" ")            
            if parentheticals[0] != 'NONE':
                temp={}
                temp['sentence'] = " ".join(parentheticals)
                temp['scene_no'] = scene_no
                temp['type'] = 'DL_PARENTH'
                temp['type_no'] = ''
                temp['phrases'] = []
                temp['sentence_no_in_scene'] = 0
                temp['paragraph_no'] = 0
                temp['importance'] = 0
                temp['final_importance'] = 9999
                temp['zero_one'] = '1'
                sentences.append(temp)
            
            dialogues = par_dia[0][1].split(" ")         # dialogues
            sent_text = (par_dia[0][1]).split(".")
            temp = {}
            temp['sentence'] = " ".join(dialogues)
            temp['scene_no'] = scene_no
            temp['type'] = 'DL_DELIVERY'
            temp['type_no'] = ''
            temp['phrases'] = []
            temp['sentence_no_in_scene'] = 0
            temp['paragraph_no'] = 0
            temp['importance'] = 0
            temp['final_importance'] = 9999
            temp['zero_one'] = '1'
            sentences.append(temp)   
    scene_no = scene_no +1

# scenes list
scenes_list = []
val = 1
scene_no = 0   
for scene in scenes:
    if scene_no == 0:
        scene_no = scene_no+1
        continue
    paragraph_no = 0
    temp = {}
    temp['scene_no'] = scene_no
    for line in scene:
        if type(line) == type(""):                        
            if line.startswith('INT') or line.startswith('EXT') or line.startswith('EXT/INT') or line.startswith('INT/EXT'):
                temp['slug line'] = line
            else: 
                paragraph_no += 1     #AC
        else:
            pass
    temp['importance'] = val
    scenes_list.append(temp)
    scene_no = scene_no+1
    val += 1

# character importanc will be used mainly with dialogues
characters_importance = []  
for character in characters:
    temp = {}
    temp['character'] = character
    temp['importance'] = character_importance(character)
    temp['relative_importance'] = 0
    characters_importance.append(temp)

# # sort character importance
characters_importance = sorted(characters_importance, key = lambda k: k['importance'])

val = 1
for char in characters_importance:
    ch = char['character']
for character_dict in characters_importance:
    if character_dict['character'].lower() == ch.lower():
        val -= 1 
        character_dict['relative_importance'] += val    
        ch = character_dict['character']
    else: 
        character_dict['relative_importance'] += val 
        ch = character_dict['character']
    val += 1 

# relative importance of a phrase in the script
# find phrases
for each_sentence in sentences:
    if each_sentence['type'] == 'AC':
        find_phrases(each_sentence['sentence'], each_sentence['scene_no'], each_sentence['sentence_no_in_scene'], each_sentence['type'], each_sentence['paragraph_no'])

# assign inherent importance to phrases
VP = 5
NP = 4
for each_phrase in phrases:
    if each_phrase['type'] == 'AC':
        if each_phrase['phrase_type'] == 'VP': 
            each_phrase['importance'] = each_phrase['importance'] * VP
        elif each_phrase['phrase_type']== 'NP':
            each_phrase['importance'] = each_phrase['importance'] * NP  
        else:
            each_phrase['importance'] = each_phrase['importance'] * 2

# # # importance of sentence in paragraph 
for each_sentence in sentences:
    if each_sentence['type'] == 'AC':
        for each_dict in each_sentence['phrases']:
            phrase = each_dict['phrase']
            set_importance(phrase, each_sentence['sentence_no_in_scene'], each_sentence['scene_no'], each_sentence['importance'])

# Note: we are updating importance of phrases in the phrases list, not in the sentences list

# relative importance of paragraph in scene and scene in script 
for each_phrase in phrases:
    if each_phrase['type'] == 'AC':
        each_phrase['importance'] = each_phrase['importance'] * each_phrase['paragraph_no'] * each_phrase['scene_no']

# add relative importance of characters in the phrases
# eg: if phrase == 'KUSH':
#        phrase['importance'] = phrase['importance'] * ( relative importance of KUSH )
for each_phrase in phrases:
    if each_phrase['type'] == 'AC':
        ph = each_phrase['phrase'].upper()
        if ph.strip(' ') in characters:
            for ch in characters_importance:
                if ch['character'] == ph.strip(' '):
                    importance = ch['relative_importance']
                    set_importance1(ph, each_phrase['sentence_no_in_scene'], each_phrase['scene_no'], importance)

# calculating sentence importance by adding the importance of the phrases in it
# Eg: sentence['importance'] = phrase1 + phrase2 + phrase3...
for each_sentence in sentences:
    if each_sentence['type'] == 'AC':
        sentence = each_sentence['sentence']
        for each_dict in each_sentence['phrases']:
            ph = each_dict['phrase']
            imp = set_importance2(ph, each_sentence['scene_no'], each_sentence['sentence_no_in_scene'])
            each_sentence['final_importance'] += imp

# for p in sentences:
#     print(p)
#     print()

################################ Integrating

def find_threshold(retain_percent):
#     """ Takes as input the percentage of compression required in the script 
#     and returns the 'threshold_val' - the thershold value of IMPORTANCE 
#                 the 'threshold_count' - the number of words in the reduced script
#     """
    number_of_sents = 0 
    sent_temp = []
    for each_sentence in sentences:
        if each_sentence['type'] != 'DL_SPEAKER' and each_sentence['type'] != 'DL_PARENTH' and each_sentence['type'] != 'DL_DELIVERY':
            sent_temp.append(each_sentence['final_importance'])
            number_of_sents += 1   
        #print (word_temp[number_of_words-1],number_of_words)   
    sent_temp.sort() 
    print("importance: ")
    print(sent_temp)
    reduced_percent = 1 - retain_percent/100   #output script
    print("redu: ",reduced_percent)
    threshold_count = round(number_of_sents * reduced_percent)
    threshold_val = sent_temp[threshold_count]
    return threshold_val, threshold_count

def find_threshold_removal_impact(retain_percent):
    """ Takes as input the percentage of compression required in the script 
    and returns the 'threshold_val' - the thershold value of REMOVAL IMPACT 
                the 'threshold_count' - the number of words in the reduced script
    """
    number_of_sents = 0
    sent_temp=[]
    for each_sentence  in sentences:
        if each_sentence['type'] != 'DL_SPEAKER' and each_sentence['type'] != 'DL_PARENTH' and each_sentence['type'] != 'DL_DELIVERY':
            sent_temp.append(each_sentence['removal_impact'])
            number_of_sents += 1   
        #print (word_temp[number_of_words-1],number_of_words)   
    sent_temp.sort()
    print() 
    print("Removal impact: ")
    print(sent_temp)
    retain_percent = 1 - retain_percent/100
    #print("redu=",reduced_percent)
    threshold_count = round(number_of_sents * retain_percent)
    threshold_val = sent_temp[threshold_count]
    return threshold_val, threshold_count

def set_zero_initial(retain_percent):
#     """ Takes as input the percentage of reduced script and sets the words as '0' and '1' based on 
#         the threshold value of importance
#     """
    threshold_v, threshold_c = find_threshold(retain_percent) #find threshold
    # sentence_no = 1
    cnt=0
    non_rem_cnt=0
    print()
    print ("threshold value ", threshold_v)
    print()
    for each_sentence in sentences:
        if each_sentence['final_importance'] < threshold_v:
            each_sentence['zero_one'] = '0'     #remove
            cnt += 1
        else:
            each_sentence['zero_one'] = '1'
            non_rem_cnt += 1
        sent_importance.append(each_sentence['final_importance'])
        #print("importance zero=", each_word['zero_one'])
    print (" Number of sentences below threshold value", cnt)
    print (" Number of sentences above threshold value", non_rem_cnt)
    return threshold_c

def set_zero_initial_removal_impact(retain_percent):
    """ Takes as input the percentage of reduced script and sets the words as '0' and '1' based on 
        the threshold value of removal impact
    """
    threshold_v, threshold_c = find_threshold_removal_impact(retain_percent) 
    for each_sentence in sentences:
        if each_sentence['removal_impact'] <= threshold_v:
            each_sentence['zero_one'] = '0'  #removable  
        else:
            each_sentence['zero_one'] = '1'
    return threshold_c

def convert_importance_to_priority(sent_importance):
#     """
# #     Takes the 'word_importance' as the input and define the 'priority' of the word for removal based on WORD IMPORTANCE
# #     e.g. if priority[1] = j   this means that the first priority for removal is of the  'jth WORD' 
#             priority[5] = 7   means 5th priority for removal is of the 7th word
#             CAN BE MODIFIED USING SORTING OF REDUCED WORD IMPORTANCE
#     """
    reduced_sent_imp = list(sent_importance)
    length = len(sent_importance)
    minpos = sent_importance.index(min(sent_importance))
    #minpos is the index of the min value in the word_importance list
    k = 0
    for i in range(length):
        priority.append(0)    #set all values as 0

    while(k < length):
        min_value = min(reduced_sent_imp)          #reduced_word_imp is word_importance
        minpos = sent_importance.index(min_value)
        priority[k] = minpos                      # first value in priority list will be the index of the min value in word_importance list
        reduced_sent_imp.remove(min_value)        # that word is removed
        k = k + 1

def convert_removal_impact_to_priority(removal_impact_list):
    """
    # Takes the 'removal impact list' as the input and define the 'priority' of the word for removal based on REMOVAL IMPACT VALUE OF THE WORD
    # e.g. if priority[1]=j   this means that the first priority is of the  'jth WORD' (based on removal impact of word j)
    #         priority[5]=7   means that the 5th priority for removal is of the 7th word (based on removal impact of word 7)
#     # """
    unique_removal_impact_list = list(dict.fromkeys(removal_impact_list))
    unique_removal_impact_list.sort()
    length = len(removal_impact_list)
    priority_removal=[]
    for i in range(length):
        priority_removal.append(0)
    priority_counter=0
    for each_unique in unique_removal_impact_list:
        sent_no = 0
        for each_removal_impact_value in removal_impact_list:
            if (each_unique == each_removal_impact_value):
                priority_removal[priority_counter] = sent_no
                # if priority[1]=j   this means that the first priority is of the  'jth WORD' (based on removal impact of word j)
                # priority[5]=7
                priority_counter = priority_counter + 1
            sent_no = sent_no + 1
    #print("removal priority=",priority_removal)            
    return priority_removal

def removal_impact_sentence(each_sentence):
    #  Takes the word as input and calculates the removal impact, if the sentence it contains needs to be removed
    #     Values returned are
    #     'word_removal_impact' - sum of importance of all the words in the sentence that contains it
    #     'num_impacted'        - number of words in the sentence
    #     'impacted'            - list of word_no of the words in the the sentence
    # """
    #print("word to check in sentence =",word['word'])
    comp_scene_no = each_sentence['scene_no']
    comp_scene_type = each_sentence['type']
    comp_sentence_no = each_sentence['sentence_no_in_scene']
    sent_no = 0
    num_impacted = 0
    impacted = []
    sent_removal_impact = 0
    for each_sentence in sentences:
        if(each_sentence['scene_no'] == comp_scene_no and each_sentence['type'] == comp_scene_type and each_sentence['sentence_no_in_scene'] == comp_sentence_no):
            sent_removal_impact +=  each_sentence['final_importance']
            impacted.append(sent_no)
            num_impacted = num_impacted + 1
        sent_no = sent_no + 1
    return sent_removal_impact, num_impacted, impacted      

def removal_impact_scene(each_sentence):
    """Takes the word as input and calculates the removal impact, if the SCENE it contains needs to be removed
        Values returned are
        'word_removal_impact' - sum of importance of all the words in the SCENE that contains it
        'num_impacted'        - number of words in the SCENE
        'impacted'            - list of word_no of the words in the SCENE
    """
    #print("removal of word in  scene=",word['word'])
    comp_scene_no = each_sentence['scene_no']    # compute scene number
    impacted = []                 #list of word_no of the words in the SCENE              
    num_impacted = 0
    sent_no = 0
    scene_removal_impact = 0
    for each_sentence in sentences:
        if(each_sentence['scene_no'] == comp_scene_no):   #loop will execute for all the words in that scene 
            scene_removal_impact = scene_removal_impact + each_sentence['final_importance']   #the scene removal impact will b the sum of the word importance of all the words in that scene
            impacted.append(sent_no)
            num_impacted = num_impacted + 1
        sent_no += 1
    return scene_removal_impact, num_impacted, impacted 

def removal_impact_remaining_sents(each_sentence):
    comp_scene_no = each_sentence['scene_no']
    comp_scene_type = each_sentence['type']
    sent_no = 0
    num_impacted = 0
    impacted = []
    sent_removal_impact = 0
    for each_sentence in sentences:
        if(each_sentence['scene_no'] == comp_scene_no and each_sentence['type'] == comp_scene_type):
            sent_removal_impact +=  999
            impacted.append(sent_no)
            num_impacted = num_impacted + 1
        sent_no = sent_no + 1
    return sent_removal_impact, num_impacted, impacted  


def assign_sent_removal_impact():
    """
    Calculates the removal impact of each word in the script 
        If the word is part of a slugline, the removal impact is sum of importance of all the words in the scene
        If the word is a speaker, the removal impact is sum of importance of all words in the dialoge, parenthetical and speaker
        If the word is a verb, the removal impact is sum of all words in the sentence that contains it
        DEFAULT REMOVAL IMPACT CHANGED FROM ZERO TO 1
    """
    removal_impact_value_list=[]
    for each_sentence in sentences:
        impacted_sents = []   #list of phrases impacted by the removal of the each_phrase
        each_sentence['removal_impact'] = 1
        each_sentence['impacted_sent_list'] = impacted_sents

        if(each_sentence['type'] =='SL'):   #slugline
            removal_impact, impacted_no, impacted_sents = removal_impact_scene(each_sentence)
            each_sentence['removal_impact'] = removal_impact
            each_sentence['impacted_sent_list'] = impacted_sents

        elif(each_sentence['type'] =='AC'):  
            removal_impact, impacted_no, impacted_sents = removal_impact_sentence(each_sentence)
            each_sentence['removal_impact'] = removal_impact
            each_sentence['impacted_sent_list'] = impacted_sents
        else:
            removal_impact, impacted_no, impacted_sents = removal_impact_remaining_sents(each_sentence)
            each_sentence['removal_impact'] = removal_impact
            each_sentence['impacted_word_list'] = impacted_sents

        removal_impact_value_list.append(each_sentence['removal_impact'])
    return removal_impact_value_list

retain_percent = 80
threshold_counter = set_zero_initial(retain_percent)   
convert_importance_to_priority(sent_importance)
g_removal_impact_value_list = assign_sent_removal_impact()   #calculate word impact of each word
threshold_counter_removal_impact = set_zero_initial_removal_impact(retain_percent)
r_priority = convert_removal_impact_to_priority(g_removal_impact_value_list)
# # # a_by_b_threshold = calculate_a_by_b()

for w in sentences:
    new_dict = {'zero_one': '1'}
    w.update(new_dict)

# print("threshold_counter_removal_impact: ", threshold_counter_removal_impact)
count = 0
for i in r_priority[ : threshold_counter_removal_impact]:
    for each_sentence in sentences:
        if i == sentences.index(each_sentence):
            count += 1
            new_dict = {'zero_one': '0'}
            each_sentence.update(new_dict)              # to be removed
# print("Found: ", count)
# ##########################################################################################33
# generate script
doc = docx.Document() 
style = doc.styles['Normal']
font = style.font
font.name = 'Courier New'
font.size = Pt(12)

def print_speaker(dialogue_speaker):
    spk = doc.add_paragraph(dialogue_speaker)
    spk_format = spk.paragraph_format
    spk_format.space_after = 0
    spk_format.line_spacing = Pt(12)
    spk_format.left_indent = Inches(2.5)

def print_dialogue(dialogue_delivery):
    d = doc.add_paragraph()
    d.add_run(dialogue_delivery)
    d_format = d.paragraph_format
    d_format.keep_together = True
    d_format.space_after = Pt(12)
    d_format.line_spacing = Pt(12)
    d_format.left_indent = Inches(1.5)
    d_format.right_indent = Inches(1.25)

def create_script(doc, dialogue_speaker, dialogue_delivery, dialogue_parenthetical, slugline_list, action_list): 
 
    slugline_data = " ".join(slugline_list)
    slug_data = doc.add_paragraph()
    slug_data_format = slug_data.paragraph_format   #styling
    slug_data_format.space_after = Pt(12)
    slug_data_format.keep_with_next = True   #next line
    slug_data_format.left_indent = Inches(-0.5)
    slug_data.add_run(slugline_data).bold = True
 
#action line
    action_data = " ".join(action_list)
    act = doc.add_paragraph()
    act.add_run(action_data)
    act_format = act.paragraph_format
    act_format.space_after = Pt(12)
    act_format.line_spacing = Pt(12)
    act_format.left_indent = Inches(0.5)

    for i in range(len(dialogue_speaker)):
        print_speaker(dialogue_speaker[i])
        print_dialogue(dialogue_delivery[i])

    doc.save('changed_threshold.docx')     

scene_no = len(scenes)
for s in range(1, scene_no):
    dialogue_counter = 1
    action_list = []                            # list of words in the action of the particular scene
    slugline_list = []                          # list of words in the slugline of the particular scene                               
    dialogue_speaker = []                       # speakers in the scene
    dialogue_delivery = []                      # list of words in the dialog of the particular scene
    dialogue_parenthetical = []  
    # print("Scene taken is: ", s)
    cnt_total = 0
    cnt_removed = 0
    for w in sentences:
        if w['scene_no'] == s:
            cnt_total += 1
            if w['type'] == 'SL':                     #SLUGLINE
                if w['zero_one'] == '0':   #to be removed
                    cnt_removed += 1
                    res = ''
                    for t in w['sentence']:
                        res = res + t +  '\u0336'   #strike through
                    slugline_list.append(res)
                else:
                    slugline_list.append(w['sentence'])

            elif w['type'] == 'AC':   
                if w['zero_one'] == '0':   #to be removed
                    cnt_removed += 1
                    res = ''
                    for t in w['sentence']:
                        res = res + t +  '\u0336'   #strike through
                    action_list.append(res)
                else:
                    action_list.append(w['sentence'])
            elif w['type'] == 'DL_SPEAKER':   
                    dialogue_speaker.append(w['sentence'])
            elif w['type'] == 'DL_PARENTH':   
                    dialogue_parenthetical.append(w['sentence'])
            elif w['type'] == 'DL_DELIVERY':   
                    dialogue_delivery.append(w['sentence'])

    print("Total sentence in ", s," scene: ", cnt_total)
    print("Sentences removed from",s," scene : ", cnt_removed)
    create_script(doc, dialogue_speaker, dialogue_delivery, dialogue_parenthetical, slugline_list, action_list)
print("Done!")


